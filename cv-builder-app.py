#importing a document 
#pip3 uninstall python-docx
#install package by using "pip3 install python-docx"
from docx import Document
from docx.shared import Inches
#importing python text to speach version 3
# pip3 uninstall python-docx
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()


#Adding profile picture
document.add_picture(
    '005.jpg', 
width=Inches(1.0)
)
#name, phone, email details
name = input('what is your name? ')
speak('Hello' + name + ' how are you today')

speak('what is your phone number?')
phone_number = input('what is your phone number? ')
email = input('what is your email? ')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email)

# About me
document.add_heading('about me')
about_me = input('Tell about yourself? ')
document.add_paragraph(about_me)

#work experience
document.add_heading('Work Experience')
p = document.add_paragraph('')

company = input('Enter Company ')
from_date = input('From Date ')
to_date = input('To Date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' +  to_date  + '\n').italic = True
experience_details = input(
    'Describe your experience at' + company+ ' ' )
p.add_run(experience_details)

#more experiences

while True:
    has_more_experiences = input(
        'Do you have more experiences? Yes or No ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph('')

        company = input('Enter Company ')
        from_date = input('From Date ')
        to_date = input('To Date ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' +  to_date  + '\n').italic = True
        experience_details = input(
            'Describe your experience at' + company)
        p.add_run(experience_details)
    else:
        break 

#Skills
document.add_heading('Skills')
skill = input('Enter skill ' )
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have more skills? Yes or No ')
    if has_more_skills.lower() == 'yes':
        skill = input('Enter skill')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

#footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = " CV gemerated using Amigoscode help"

document.save('cv.docx')







